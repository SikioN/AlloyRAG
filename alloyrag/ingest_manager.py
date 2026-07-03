import os
import json
import hashlib
import time
import httpx
from dotenv import load_dotenv
from alloyrag.rag_system import get_rag_instance
from alloyrag.ingest_excel import extract_excel_to_kg
from alloyrag.utils import logger

load_dotenv()

STATUS_FILE = "./rag_storage/ingested_files.json"


def calculate_md5(filepath: str) -> str:
    """Calculate MD5 hash of a file."""
    hash_md5 = hashlib.md5()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def get_ingested_states() -> dict:
    """Load states of previously ingested files."""
    if os.path.exists(STATUS_FILE):
        try:
            with open(STATUS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Failed to load status file: {e}. Starting fresh.")
    return {}


def save_ingested_states(states: dict):
    """Save states of ingested files."""
    os.makedirs(os.path.dirname(STATUS_FILE), exist_ok=True)
    try:
        with open(STATUS_FILE, "w", encoding="utf-8") as f:
            json.dump(states, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"Failed to save status file: {e}")


def check_docling_availability() -> bool:
    """Check if docling-serve is running at DOCLING_ENDPOINT."""
    endpoint = os.getenv("DOCLING_ENDPOINT", "http://localhost:5001")
    try:
        response = httpx.get(f"{endpoint}/docs", timeout=3)
        return response.status_code == 200
    except Exception:
        return False


def parse_pdf_via_docling(filepath: str) -> str | None:
    """Send PDF to docling-serve and return markdown content."""
    endpoint = os.getenv("DOCLING_ENDPOINT", "http://localhost:5001")
    url = f"{endpoint}/v1/convert/file"
    logger.info(f"Parsing PDF '{filepath}' via Docling Serve ({url})...")

    try:
        with open(filepath, "rb") as f:
            files = {"files": (os.path.basename(filepath), f, "application/pdf")}
            # Request markdown format specifically
            options = {"to_formats": ["markdown"], "do_ocr": True}
            data = {"options": json.dumps(options)}

            response = httpx.post(url, files=files, data=data, timeout=120)
            if response.status_code == 200:
                res_json = response.json()
                doc = res_json.get("document", {})
                markdown = (
                    doc.get("md_content") or doc.get("text") or res_json.get("text")
                )
                if markdown:
                    logger.info("Successfully converted PDF to Markdown via Docling.")
                    return markdown
            logger.warning(
                f"Docling Serve returned code {response.status_code}: {response.text}"
            )
    except Exception as e:
        logger.error(f"Failed to parse via Docling: {e}")
    return None


def parse_pdf_via_pypdf(filepath: str) -> str:
    """Fallback parser using standard pypdf library."""
    logger.info(f"Parsing PDF '{filepath}' locally via pypdf...")
    try:
        import pypdf

        reader = pypdf.PdfReader(filepath)
        text_parts = []
        for idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to parse via pypdf: {e}")
        return ""


def parse_docx(filepath: str) -> str:
    """Parse Word document locally."""
    logger.info(f"Parsing Word document '{filepath}' locally...")
    try:
        import docx

        doc = docx.Document(filepath)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also parse tables if present
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_cells))
            if table_text:
                text_parts.append("\n" + "\n".join(table_text) + "\n")

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to parse Word document: {e}")
        return ""


async def aingest_file(filepath: str) -> bool:
    """Route file to appropriate parser and insert into AlloyRAG asynchronously."""
    import asyncio

    ext = os.path.splitext(filepath)[1].lower()
    rag = get_rag_instance()
    await rag.initialize_storages()

    # 1. Route Excel files to custom KG pipeline
    if ext in [".xlsx", ".xls"]:
        loop = asyncio.get_running_loop()
        custom_kg = await loop.run_in_executor(None, extract_excel_to_kg, filepath)
        doc_id = f"excel_{os.path.basename(filepath)}"
        if custom_kg is None:
            import time

            now = int(time.time())
            await rag.doc_status.upsert(
                {
                    doc_id: {
                        "status": "failed",
                        "content_summary": "Excel file parsing failed (returned None)",
                        "content_length": 0,
                        "created_at": now,
                        "updated_at": now,
                        "track_id": doc_id,
                        "chunks_count": 0,
                        "file_path": filepath,
                        "error_msg": "Excel extractor returned None",
                        "metadata": {"type": "excel"},
                    }
                }
            )
            return False

        logger.info(
            f"Ingesting {len(custom_kg['entities'])} entities and {len(custom_kg['relationships'])} relations asynchronously into AlloyRAG..."
        )
        try:
            await rag.ainsert_custom_kg(custom_kg, full_doc_id=doc_id)
            import time

            now = int(time.time())
            await rag.doc_status.upsert(
                {
                    doc_id: {
                        "status": "processed",
                        "content_summary": f"Excel file: {len(custom_kg.get('entities', []))} entities, {len(custom_kg.get('relationships', []))} relationships",
                        "content_length": len(str(custom_kg)),
                        "created_at": now,
                        "updated_at": now,
                        "track_id": doc_id,
                        "chunks_count": len(custom_kg.get("chunks", [])),
                        "file_path": filepath,
                        "metadata": {"type": "excel"},
                    }
                }
            )
            logger.info(f"Successfully finished Excel ingestion for {filepath}.")
            return True
        except Exception as e:
            import time

            now = int(time.time())
            await rag.doc_status.upsert(
                {
                    doc_id: {
                        "status": "failed",
                        "content_summary": f"Excel file ingestion failed: {e}",
                        "content_length": 0,
                        "created_at": now,
                        "updated_at": now,
                        "track_id": doc_id,
                        "chunks_count": 0,
                        "file_path": filepath,
                        "error_msg": str(e),
                        "metadata": {"type": "excel"},
                    }
                }
            )
            raise

    # 2. Extract Text Content
    content = ""
    loop = asyncio.get_running_loop()
    if ext == ".pdf":
        if check_docling_availability():
            content = await loop.run_in_executor(None, parse_pdf_via_docling, filepath)

        if not content:
            logger.warning(
                "Docling parser unavailable or failed. Falling back to local pypdf extraction."
            )
            content = await loop.run_in_executor(None, parse_pdf_via_pypdf, filepath)

    elif ext in [".docx", ".doc"]:
        content = await loop.run_in_executor(None, parse_docx, filepath)

    elif ext in [".txt", ".md"]:
        try:

            def read_txt(fp):
                with open(fp, "r", encoding="utf-8") as f:
                    return f.read()

            content = await loop.run_in_executor(None, read_txt, filepath)
        except Exception as e:
            logger.error(f"Failed to read text file {filepath}: {e}")
            return False
    else:
        logger.warning(f"Unsupported file format: {ext} for {filepath}")
        return False

    if not content or not content.strip():
        logger.warning(f"No content extracted from {filepath}. Skipping.")
        return False

    # 3. Ingest into AlloyRAG asynchronously
    doc_id = os.path.basename(filepath)
    logger.info(f"Inserting document {doc_id} into AlloyRAG...")
    await rag.ainsert(content, ids=doc_id, file_paths=filepath)
    logger.info(f"Successfully indexed document {doc_id}.")
    return True


async def ascan_and_ingest_inputs():
    """Scan input folder and process new/changed files incrementally asynchronously."""
    input_dir = os.getenv("INPUT_DIR", "./inputs")
    os.makedirs(input_dir, exist_ok=True)

    logger.info(f"Scanning directory '{input_dir}' for documents...")
    states = get_ingested_states()

    files_in_dir = []
    for root, _, files in os.walk(input_dir):
        for name in files:
            # Skip hidden files
            if name.startswith("."):
                continue
            files_in_dir.append(os.path.join(root, name))

    if not files_in_dir:
        logger.info("No files found in inputs directory.")
        return

    logger.info(f"Found {len(files_in_dir)} file(s) in inputs.")
    rag = get_rag_instance()
    await rag.initialize_storages()
    states_changed = False

    for filepath in files_in_dir:
        filename = os.path.basename(filepath)
        current_md5 = calculate_md5(filepath)
        mtime = os.path.getmtime(filepath)

        # Check if already processed and exists in DB
        state = states.get(filepath)
        doc_id = os.path.basename(filepath)
        db_doc = await rag.doc_status.get_by_id(doc_id)
        
        db_status = ""
        if db_doc:
            db_status_raw = db_doc.get("status", "")
            if hasattr(db_status_raw, "value"):
                db_status = db_status_raw.value
            else:
                db_status = str(db_status_raw)

        if (
            state
            and state.get("md5") == current_md5
            and state.get("status") == "success"
            and db_status == "processed"
        ):
            logger.info(f"File '{filename}' is unchanged and successfully processed. Skipping.")
            continue

        logger.info(f"Processing new or changed file: '{filename}'...")
        success = False
        try:
            success = await aingest_file(filepath)
        except Exception as e:
            logger.error(f"Exception during ingestion of {filename}: {e}")

        states[filepath] = {
            "md5": current_md5,
            "mtime": mtime,
            "status": "success" if success else "failed",
            "processed_at": time.time(),
        }
        states_changed = True

    if states_changed:
        save_ingested_states(states)
        logger.info("Incremental state index updated.")


def scan_and_ingest_inputs():
    """Scan input folder synchronously (wrapper)."""
    import asyncio

    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        loop = None

    if loop and loop.is_running():
        import concurrent.futures

        with concurrent.futures.ThreadPoolExecutor() as pool:
            return pool.submit(lambda: asyncio.run(ascan_and_ingest_inputs())).result()
    else:
        return asyncio.run(ascan_and_ingest_inputs())


if __name__ == "__main__":
    scan_and_ingest_inputs()
